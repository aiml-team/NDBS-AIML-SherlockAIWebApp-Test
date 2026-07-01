from docxtpl import DocxTemplate, InlineImage
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
from io import BytesIO
import base64
import re
import os
import json


# ── Internet Research formatting ─────────────────────────────────────────────
# The Tavily enrichment block (sherlock Web App/tavily_search.py) embeds plain
# text like:
#     [Internet Research]
#     • <fact bullet>
#     • Source: <title> — <url>
# into a field's content. docxtpl renders that as a single paragraph with soft
# line breaks, so we post-process the rendered DOCX to:
#   • restyle the "[Internet Research]" header line as italic + red
#   • restyle each "• Source:" line as 8 pt + grey
_HEADER_TEXT = "[Internet Research]"
_RED_HEX = "C00000"   # MS Office "Dark Red"
_GREY_HEX = "808080"  # 50% grey
_SOURCE_SZ_HALF_POINTS = "16"  # 16 half-points = 8 pt


def _classify_line(text):
    s = (text or "").strip()
    if s == _HEADER_TEXT:
        return "header"
    if s.startswith("• Source:") or s.startswith("Source:"):
        return "source"
    return "normal"


def _iter_all_paragraphs(doc):
    """Yield every w:p paragraph in the document, including those inside
    (possibly nested) tables."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _collect_lines(p_el):
    """Return [(text, base_rPr_or_None), ...] for the paragraph, splitting on
    <w:br/> elements wherever they appear. base_rPr is the first non-empty
    run-properties element seen for that line — we deepcopy it when rebuilding
    so the original run styling (font, bold, etc.) is preserved for normal
    lines."""
    lines = []
    cur_text = ""
    cur_rpr = None
    for r in p_el.findall(qn("w:r")):
        r_rpr = r.find(qn("w:rPr"))
        for child in list(r):
            tag = child.tag
            if tag == qn("w:t"):
                if cur_rpr is None and r_rpr is not None:
                    cur_rpr = r_rpr
                cur_text += (child.text or "")
            elif tag == qn("w:br"):
                lines.append((cur_text, cur_rpr))
                cur_text = ""
                cur_rpr = None
            elif tag == qn("w:tab"):
                cur_text += "\t"
    lines.append((cur_text, cur_rpr))
    return lines


def _apply_header_style(rpr):
    for tag in ("w:i", "w:iCs", "w:color"):
        for el in rpr.findall(qn(tag)):
            rpr.remove(el)
    rpr.append(OxmlElement("w:i"))
    rpr.append(OxmlElement("w:iCs"))
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _RED_HEX)
    rpr.append(color)


def _apply_source_style(rpr):
    for tag in ("w:color", "w:sz", "w:szCs"):
        for el in rpr.findall(qn(tag)):
            rpr.remove(el)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _GREY_HEX)
    rpr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), _SOURCE_SZ_HALF_POINTS)
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), _SOURCE_SZ_HALF_POINTS)
    rpr.append(szCs)


def _restyle_paragraph(p):
    p_el = p._element
    lines = _collect_lines(p_el)
    if not any(_classify_line(t) != "normal" for t, _ in lines):
        return

    # Remove existing top-level run elements; keep w:pPr and anything else.
    for r in list(p_el.findall(qn("w:r"))):
        p_el.remove(r)

    pPr = p_el.find(qn("w:pPr"))
    insert_idx = list(p_el).index(pPr) + 1 if pPr is not None else 0

    new_elems = []
    for i, (text, base_rpr) in enumerate(lines):
        if i > 0:
            br_run = OxmlElement("w:r")
            if base_rpr is not None:
                br_run.append(deepcopy(base_rpr))
            br_run.append(OxmlElement("w:br"))
            new_elems.append(br_run)
        if not text:
            continue
        cls = _classify_line(text)
        new_r = OxmlElement("w:r")
        rpr = deepcopy(base_rpr) if base_rpr is not None else OxmlElement("w:rPr")
        if cls == "header":
            _apply_header_style(rpr)
        elif cls == "source":
            _apply_source_style(rpr)
        if list(rpr):
            new_r.append(rpr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        new_r.append(t)
        new_elems.append(new_r)

    for el in new_elems:
        p_el.insert(insert_idx, el)
        insert_idx += 1


def _restyle_internet_research(doc):
    """Walk the document and restyle any [Internet Research] header / Source
    bullet lines that ended up inside a docxtpl-rendered paragraph."""
    for p in _iter_all_paragraphs(doc):
        _restyle_paragraph(p)

def process_json_data_for_template(json_data):
    """
    Process the nested JSON data structure and extract content with image placeholders
    """
    processed_data = {}
    all_images = {}
    
    for main_section, main_content in json_data.items():
        print(f"Processing main section: {main_section}")
        
        if not isinstance(main_content, dict):
            continue
        
        for sub_section, sub_content in main_content.items():
            print(f"  Processing sub-section: {sub_section}")
            
            # Create a unique key for this section
            section_key = f"{main_section}_{sub_section}".replace(" ", "_").replace("/", "_").replace("(", "").replace(")", "").replace("&", "and").replace('"', "").replace("'", "")
            
            # Get content and images
            if isinstance(sub_content, dict):
                content = sub_content.get('content', '')
                images = sub_content.get('images', {})
            else:
                content = str(sub_content)
                images = {}
            
            # Store processed content
            processed_data[section_key] = content
            
            # Collect all images with their original keys
            for img_key, img_data in images.items():
                all_images[img_key] = img_data
                print(f"    Found image: {img_key}")
    
    return processed_data, all_images

def sanitize_key(key_string):
    """
    Sanitize a key string by removing/replacing all problematic characters
    """
    safe_key = key_string
    
    # Replace specific characters one by one
    safe_key = safe_key.replace('"', '')
    safe_key = safe_key.replace("'", '')
    safe_key = safe_key.replace('&', 'and')
    safe_key = safe_key.replace(' ', '_')
    safe_key = safe_key.replace('(', '')
    safe_key = safe_key.replace(')', '')
    safe_key = safe_key.replace('/', '___')
    safe_key = safe_key.replace(';', '_')
    safe_key = safe_key.replace(':', '')
    safe_key = safe_key.replace(',', '')
    safe_key = safe_key.replace('-', '_')
    
    # Remove any remaining special characters using regex
    safe_key = re.sub(r'[^\w_]', '', safe_key)
    
    # Clean up multiple underscores
    safe_key = re.sub(r'_+', '_', safe_key)
    safe_key = safe_key.strip('_')
    
    return safe_key

def create_safe_key_mapping(json_data):
    """
    Create safe template keys and a mapping to avoid special character issues
    """
    safe_mapping = {}
    key_to_safe = {}
    
    for main_section, main_content in json_data.items():
        # Create safe main section key
        safe_main = sanitize_key(main_section)
        
        safe_mapping[safe_main] = {}
        key_to_safe[main_section] = safe_main
        
        print(f"  Main section mapping: '{main_section}' → '{safe_main}'")
        
        if not isinstance(main_content, dict):
            continue
            
        for sub_section, sub_content in main_content.items():
            # Create safe sub section key
            safe_sub = sanitize_key(sub_section)
            
            safe_mapping[safe_main][safe_sub] = sub_content
            key_to_safe[f"{main_section}|{sub_section}"] = f"{safe_main}|{safe_sub}"
            
            print(f"    Sub-section mapping: '{sub_section}' → '{safe_sub}'")
    
    return safe_mapping, key_to_safe

def render_step1_json_text_conversion_memory(template_path, json_data, prospect_name="",
                                              internet_research_used=False):
    """
    STEP 1: Convert [IMAGE_1] → {{image_1}} and render text with JSON data.

    `internet_research_used` controls the Jinja conditional under the
    'Customer/Business Overview' heading: when True, a disclaimer paragraph
    informing the reader that Section 1 was enriched from the internet is
    rendered; when False, the entire paragraph is removed.

    Returns: BytesIO object containing intermediate document
    """
    try:
        print("  Loading original template...")
        doc = DocxTemplate(template_path)
        
        # Create safe key mapping to avoid special character issues
        safe_data, key_mapping = create_safe_key_mapping(json_data)
        
        # Process content and convert image placeholders
        for main_key, main_content in safe_data.items():
            print(f"  Processing main section: {main_key}")
            
            if not isinstance(main_content, dict):
                continue
            
            for sub_key, sub_content in main_content.items():
                print(f"    Processing sub-section: {sub_key}")
                
                if isinstance(sub_content, dict):
                    content = sub_content.get('content', '')
                    images = sub_content.get('images', {})
                else:
                    content = str(sub_content)
                    images = {}
                
                # Convert image placeholders to template variables in the content
                modified_content = content
                # Updated regex to match IMAGE_XXXXX (alphanumeric hash and underscore) format
                # Also handle potential spacing issues
                image_placeholders = re.findall(r'\[\s*IMAGE_[A-Z0-9]+\s*\]', content)
                
                print(f"      Original content preview: {content[:200]}")
                print(f"      Content type: {type(content)}, length: {len(content)}")
                print(f"      Found {len(image_placeholders)} image placeholders: {image_placeholders}")
                
                if not image_placeholders:
                    # Try alternative patterns
                    alt_matches = re.findall(r'IMAGE_[A-Z0-9]+', content)
                    if alt_matches:
                        print(f"      WARNING: Found image refs without brackets: {alt_matches}")
                
                for placeholder in image_placeholders:
                    image_key = placeholder.strip().strip('[]').strip()  # Remove brackets and whitespace
                    template_var = f'{{{{ {image_key.lower()} }}}}'  # Convert to: {{ image_a3f2b1c4d5e6f7g8 }}
                    modified_content = modified_content.replace(placeholder, template_var)
                    print(f"      Converted: '{placeholder}' → '{template_var}'")
                
                print(f"      Modified content preview: {modified_content[:200]}")
                
                # Update the content with modified version
                safe_data[main_key][sub_key] = {
                    'content': modified_content,
                    'images': images
                }
                
                if modified_content != content:
                    print(f"      Modified content: {modified_content[:100]}...")
        
        # Pre-populate ALL sections and sub-sections the template references.
        # If a key is missing from safe_data, Jinja2 raises UndefinedError on the
        # intermediate dot-access BEFORE | default('') can catch it.
        # Providing empty dicts for every possible key prevents this entirely.
        EMPTY_SUB = {
            "Current_Processes_Key_Findings":  {"content": "", "images": {}},
            "Pain_Points":                     {"content": "", "images": {}},
            "Proposed_SAP_Solutions_Mapping":  {"content": "", "images": {}},
            "Major_Gaps_and_Integrations":     {"content": "", "images": {}},
        }
        GENERAL_BUSINESS_OVERVIEW_SUBS = {
            "Areas_of_Perceived_Competitive_Advantage": {"content": "", "images": {}},
            "Business_Locations":                       {"content": "", "images": {}},
            "Contacts_Identified":                      {"content": "", "images": {}},
            "Fiscal_Year_Format":                       {"content": "", "images": {}},
            "Industry_Categorization":                  {"content": "", "images": {}},
            "Key_Public_Cloud_Disqualifiers":            {"content": "", "images": {}},
            "Key_Value_Drivers":                        {"content": "", "images": {}},
            "Legal_Entities_and_Names":                 {"content": "", "images": {}},
            "Motivations_for_Transformation":           {"content": "", "images": {}},
            "Perceived_Change_Resistance":              {"content": "", "images": {}},
            "Regulatory_Compliance_Requirements":       {"content": "", "images": {}},
            "Revenue_Band":                             {"content": "", "images": {}},
            "Schedule_of_Events":                       {"content": "", "images": {}},
            "System_Landscape":                         {"content": "", "images": {}},
            "Technical_Challenges_and_Requirements":    {"content": "", "images": {}},
            "Total_SAP_Users":                          {"content": "", "images": {}},
            "Transformation_Program_C_Suite_KPIs":      {"content": "", "images": {}},
        }
        ALL_TEMPLATE_SECTIONS = {
            "General_Business_Overview":                              GENERAL_BUSINESS_OVERVIEW_SUBS,
            "Idea_to_Market":                                         EMPTY_SUB,
            "Source_to_Pay_S2P":                                      EMPTY_SUB,
            "Plan_to_Produce_P2P":                                    EMPTY_SUB,
            "Detect_to_Correct_D2C":                                  EMPTY_SUB,
            "Forecast_to_Fulfill_F2F":                                EMPTY_SUB,
            "Warehouse_Execution_WM_EWM":                             EMPTY_SUB,
            "Lead_to_Cash_L2C":                                       EMPTY_SUB,
            "Logistics_Planning_and_Transportation_TM":               EMPTY_SUB,
            "Request_to_Service_R2S":                                 EMPTY_SUB,
            "Record_to_Report_R2R":                                   EMPTY_SUB,
            "Acquire_to_Dispose_A2D":                                 EMPTY_SUB,
            "Environmental_Social_and_Governance_ESG_Processes":      EMPTY_SUB,
            "Hire_to_Retire_H2R":                                     EMPTY_SUB,
            "Enterprise_Reporting_Data_and_Analytics_Strategy":       EMPTY_SUB,
        }
        for section_key, default_subs in ALL_TEMPLATE_SECTIONS.items():
            if section_key not in safe_data:
                safe_data[section_key] = {}
            # Fill in any missing sub-keys within each section
            for sub_key, default_val in default_subs.items():
                if sub_key not in safe_data[section_key]:
                    safe_data[section_key][sub_key] = default_val

        # Prepare context with safe keys.
        #
        # `internet_research_used` drives the {% if %} block we placed in the
        # template right after the 'Customer/Business Overview' heading. The
        # surrounding paragraph (and its Jinja tags) get removed cleanly when
        # the flag is False, so this is safe to render unconditionally.
        context = {
            'data': safe_data,
            'prospect_name': prospect_name or '',
            'internet_research_used': bool(internet_research_used),
        }

        print(f"  Context prepared with safe main sections: {list(safe_data.keys())}")
        print(f"  prospect_name in context: {prospect_name!r}")
        print(f"  internet_research_used in context: {internet_research_used!r}")
        
        # Render with text only and save to BytesIO
        doc.render(context)
        
        # Save to memory instead of file
        intermediate_bytes = BytesIO()
        doc.save(intermediate_bytes)
        intermediate_bytes.seek(0)  # Reset pointer to beginning
        
        print(f"  ✓ Step 1 completed. Intermediate document: {len(intermediate_bytes.getvalue())} bytes")
        return intermediate_bytes
            
    except Exception as e:
        print(f"  ✗ Step 1 error: {e}")
        import traceback
        traceback.print_exc()
        return None

def render_step2_with_images_json_memory(intermediate_doc_bytes, images):
    """
    STEP 2: Use intermediate document as template and render with actual images
    Returns: BytesIO object containing final document
    """
    try:
        print("  Loading intermediate template from memory...")
        
        # Load the intermediate document from BytesIO
        intermediate_doc_bytes.seek(0)  # Make sure we're at the beginning
        doc = DocxTemplate(intermediate_doc_bytes)
        
        print(f"  Processing {len(images)} images...")
        
        # Process images and create InlineImage objects
        context = {}
        
        for key, base64_data in images.items():
            try:
                print(f"    Processing {key}...")
                
                # Skip if no image data
                if not base64_data or str(base64_data).strip() == '':
                    print(f"    ⚠ {key}: No image data, skipping")
                    context[key.lower()] = f"[{key} - no data]"
                    continue
                
                # Convert to string if not already
                base64_data = str(base64_data)
                
                # Clean base64 data
                if ',' in base64_data:
                    base64_data = base64_data.split(',')[1]
                
                # Validate base64 data length
                if len(base64_data) < 100:
                    print(f"    ⚠ {key}: Base64 data too short ({len(base64_data)} chars), skipping")
                    context[key.lower()] = f"[{key} - invalid data]"
                    continue
                
                # Try to decode base64
                try:
                    image_data = base64.b64decode(base64_data)
                except Exception as decode_error:
                    print(f"    ✗ {key}: Base64 decode failed: {decode_error}")
                    context[key.lower()] = f"[{key} - decode error]"
                    continue
                
                if len(image_data) < 50:
                    print(f"    ⚠ {key}: Decoded image too small ({len(image_data)} bytes), skipping")
                    context[key.lower()] = f"[{key} - invalid image]"
                    continue
                    
                image_stream = BytesIO(image_data)
                
                # Create InlineImage object
                try:
                    image_obj = InlineImage(
                        doc, 
                        image_stream, 
                        width=Inches(2.5),
                        height=Inches(1.8)
                    )
                    
                    context[key.lower()] = image_obj
                    print(f"    ✓ {key} processed successfully")
                    
                except Exception as image_error:
                    print(f"    ✗ Error creating InlineImage for {key}: {image_error}")
                    context[key.lower()] = f"[{key} - image creation error]"
                
            except Exception as e:
                print(f"    ✗ Error processing {key}: {e}")
                context[key.lower()] = f"[{key} - error: {str(e)[:50]}]"
        
        print(f"  Context prepared with keys: {list(context.keys())}")
        
        # Render with images and save to BytesIO
        doc.render(context)

        # Post-process: restyle [Internet Research] header (italic + red) and
        # • Source: bullet lines (8 pt + grey) inserted by Tavily enrichment.
        try:
            _restyle_internet_research(doc.docx)
        except Exception as restyle_err:
            print(f"  ⚠ Internet Research restyle skipped: {restyle_err}")

        final_bytes = BytesIO()
        doc.save(final_bytes)
        final_bytes.seek(0)  # Reset pointer to beginning
        
        print(f"  ✓ Step 2 completed. Final document: {len(final_bytes.getvalue())} bytes")
        
        if len(final_bytes.getvalue()) > 10000:
            print(f"  ✓ File size suggests images were successfully embedded")
        
        return final_bytes
            
    except Exception as e:
        print(f"  ✗ Step 2 error: {e}")
        import traceback
        traceback.print_exc()
        return None

def two_step_document_creation_json_memory(template_path, json_data, prospect_name="",
                                            internet_research_used=False):
    """
    Two-step rendering process for JSON data - RETURNS BINARY DATA instead of saving
    Step 1: Process JSON data and render text with image placeholders converted to {{image_1}}, {{image_2}}
    Step 2: Use Step 1 output as new template and render with actual images
    Returns: BytesIO object containing the final Word document
    """
    
    try:
        print("=== TWO-STEP RENDERING PROCESS FOR JSON DATA (MEMORY MODE) ===")
        print(f"Template: {template_path}")
        print("Output: BytesIO object (no file saving)")
        print()
        
        # Validate template file
        if not os.path.exists(template_path):
            print(f"✗ Template file not found: {template_path}")
            return None
        
        # Process JSON data
        print("PREPROCESSING: Extracting data from JSON structure...")
        processed_sections, all_images = process_json_data_for_template(json_data)
        
        print(f"Extracted {len(processed_sections)} sections and {len(all_images)} images")
        print(f"Sections: {list(processed_sections.keys())}")
        print(f"Images: {list(all_images.keys())}")
        print()
        
        # STEP 1: Convert image placeholders and render text only
        print("STEP 1: Converting image placeholders to template variables...")
        intermediate_doc_bytes = render_step1_json_text_conversion_memory(
            template_path, json_data, prospect_name=prospect_name,
            internet_research_used=internet_research_used,
        )
        
        if not intermediate_doc_bytes:
            print("✗ Step 1 failed!")
            return None
        
        # STEP 2: Use intermediate document as template and render with images
        print("\nSTEP 2: Rendering images into the converted template...")
        final_doc_bytes = render_step2_with_images_json_memory(intermediate_doc_bytes, all_images)
        
        if not final_doc_bytes:
            print("✗ Step 2 failed!")
            return None
        
        print(f"\n🎉 SUCCESS! Two-step rendering completed!")
        print(f"Final document: BytesIO object with {len(final_doc_bytes.getvalue())} bytes")
        
        return final_doc_bytes
        
    except Exception as e:
        print(f"✗ Two-step rendering failed: {e}")
        import traceback
        traceback.print_exc()
        return None

def generate_document_in_memory(template_file, json_data, prospect_name="",
                                 internet_research_used=False):
    """
    Generate Word document and return as BytesIO object

    Args:
        template_file (str): Path to the Word template file
        json_data (dict): JSON data structure
        prospect_name (str): Prospect/client name to fill the {{ prospect_name }}
            placeholder on the cover page.
        internet_research_used (bool): Whether Tavily enriched any Section 1
            (Customer/Business Overview) field. When True, the template
            renders a fact-check disclaimer paragraph immediately under the
            Section 1 heading; when False, the disclaimer paragraph is
            removed entirely from the output.

    Returns:
        BytesIO: Document bytes or None if failed
    """
    # Validate inputs
    if not template_file or not os.path.exists(template_file):
        print(f"✗ Template file not found: {template_file}")
        return None
    
    if not json_data or not isinstance(json_data, dict):
        print("✗ Invalid JSON data provided")
        return None
    
    # Run the two-step process
    document_bytes = two_step_document_creation_json_memory(
        template_file, json_data, prospect_name=prospect_name,
        internet_research_used=internet_research_used,
    )
    
    if document_bytes:
        print(f"\nDocument generated in memory successfully!")
        print(f"Document size: {len(document_bytes.getvalue())} bytes")
        print("No files saved to disk")
        print("Document ready for download/streaming")
        
        return document_bytes
    else:
        print("\nDocument generation failed. Check the error messages above.")
        return None

def save_bytes_to_file(document_bytes, output_path):
    """
    Optional: Save the BytesIO object to a file
    
    Args:
        document_bytes (BytesIO): Document bytes
        output_path (str): Output file path
    
    Returns:
        bool: Success status
    """
    try:
        if not document_bytes:
            print("No document bytes to save")
            return False
            
        document_bytes.seek(0)  # Reset to beginning
        with open(output_path, 'wb') as f:
            f.write(document_bytes.getvalue())
        print(f"Document saved to: {output_path}")
        return True
    except Exception as e:
        print(f"Failed to save file: {e}")
        return False

# Example usage and testing
if __name__ == "__main__":
    # Test configuration
    template_file = "word_template.docx"
    
    # Test JSON data structure
    test_json_data = {
        "Enterprise Reporting; Data & Analytics Strategy": {
            "General Notes & \"Wish List\"": {
                "content": "All business units use Excel for reporting but face challenges integrating financial and operational data. [IMAGE_1] Current reporting processes are manual and time-consuming.",
                "images": {
                    "IMAGE_1": "/9j/4AAQSkZJRgABAQEAYABgAAD/2wBDAAMCAgMCAgMDAwMEAwMEBQgFBQQEBQoHBwYIDAoMDAsKCwsNDhIQDQ4RDgsLEBYQERMUFRUVDA8XGBYUGBIUFRT/2wBDAQMEBAUEBQkFBQkUDQsNFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBT/wAARCAABAAEDASIAAhEBAxEB/8QAFQABAQAAAAAAAAAAAAAAAAAAAAv/xAAUEAEAAAAAAAAAAAAAAAAAAAAA/8QAFQEBAQAAAAAAAAAAAAAAAAAAAAX/xAAUEQEAAAAAAAAAAAAAAAAAAAAA/9oADAMBAAIRAxEAPwCdABmX/9k="
                }
            },
            "Team Dynamics": {
                "content": "The analytics team consists of 5 members with varying skill levels in data analysis and reporting tools.",
                "images": {}
            }
        },
        "Record to Report (R2R)": {
            "General Ledger Accounting": {
                "content": "Current GL processes are heavily manual with monthly close taking 15 business days. Key challenges include account reconciliations and intercompany eliminations.",
                "images": {}
            }
        }
    }
    
    print("=== MEMORY-BASED DOCUMENT GENERATION TEST ===")
    print("Testing Word document generation in memory...")
    print()
    
    # Check if template exists
    if not os.path.exists(template_file):
        print(f"Warning: Template file not found: {template_file}")
        print("Please ensure the template file exists to run the test.")
    else:
        # Generate document in memory
        document_bytes = generate_document_in_memory(template_file, test_json_data)
        
        if document_bytes:
            print("\n=== SUCCESS ===")
            print("Document generated successfully in memory")
            print(f"Document size: {len(document_bytes.getvalue())} bytes")
            print("\nYou can now:")
            print("1. Return it from a web API endpoint")
            print("2. Stream it to a client")
            print("3. Save it to a file if needed:")
            print("   save_bytes_to_file(document_bytes, 'output.docx')")
            
            # Optionally save for testing
            # save_bytes_to_file(document_bytes, "test_output.docx")
        else:
            print("\n=== FAILURE ===")
            print("Document generation failed")
    
    print(f"\n=== BENEFITS FOR AZURE DEPLOYMENT ===")
    print("- No temporary files created")
    print("- Memory-based processing suitable for serverless")
    print("- Better performance and security")
    print("- Handles multiple concurrent requests")